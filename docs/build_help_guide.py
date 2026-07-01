from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Elliott_Wave_Terminal_User_Guide.docx"
NAVY = "132238"
BLUE = "1976A3"
CYAN = "00AFCB"
GREEN = "168A62"
RED = "C43D4B"
GOLD = "A97200"
LIGHT = "E8EEF5"
PALE = "F4F6F9"
INK = "1E2936"
MUTED = "647184"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.style = "Table Grid"
    for index, (header, width) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[index]
        cell.text = header
        set_cell_width(cell, width)
        shade(cell, LIGHT)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell_margins(cell)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(NAVY)
            run.font.size = Pt(9.5)
    header_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    header_pr.append(repeat)
    for row in rows:
        cells = table.add_row().cells
        row_pr = table.rows[-1]._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        row_pr.append(cant_split)
        for index, (value, width) in enumerate(zip(row, widths)):
            cells[index].text = str(value)
            set_cell_width(cells[index], width)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_margins(cells[index])
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(9.2)
                    run.font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_note(doc, label, text, color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_width(cell, 9360)
    shade(cell, PALE)
    cell_margins(cell, 120, 160, 120, 160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    lead = p.add_run(f"{label}: ")
    lead.bold = True
    lead.font.color.rgb = RGBColor.from_string(color)
    body = p.add_run(text)
    body.font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def step(doc, title, text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(title + ". ")
    run.bold = True
    p.add_run(text)
    return p


def page_break(doc):
    doc.add_page_break()


def add_picture(doc, path, width, alt_text):
    doc.add_picture(str(path), width=width)
    shape = doc.inline_shapes[-1]
    shape._inline.docPr.set("descr", alt_text)
    shape._inline.docPr.set("title", alt_text)
    return doc.paragraphs[-1]


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.78)
section.bottom_margin = Inches(0.72)
section.left_margin = Inches(0.82)
section.right_margin = Inches(0.82)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15
for name, size, before, after, color in (
    ("Heading 1", 16, 16, 8, BLUE),
    ("Heading 2", 13, 12, 6, BLUE),
    ("Heading 3", 11.5, 9, 4, NAVY),
):
    style = styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True
for name in ("List Bullet", "List Number"):
    styles[name].font.name = "Calibri"
    styles[name].font.size = Pt(10.3)
    styles[name].paragraph_format.space_after = Pt(4)

header = section.header
hp = header.paragraphs[0]
hp.text = "ELLIOTT WAVE TERMINAL  |  OPERATOR GUIDE"
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for run in hp.runs:
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(MUTED)
footer = section.footer
fp = footer.paragraphs[0]
fp.text = "Decision support only - validate risk, liquidity, and execution independently."
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in fp.runs:
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)

# Cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(78)
p.paragraph_format.space_after = Pt(10)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("ELLIOTT WAVE TERMINAL")
r.bold = True
r.font.name = "Calibri"
r.font.size = Pt(30)
r.font.color.rgb = RGBColor.from_string(NAVY)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(20)
r = p.add_run("User Guide and Trading Decision-Support Handbook")
r.font.size = Pt(16)
r.font.color.rgb = RGBColor.from_string(BLUE)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(28)
r = p.add_run("Deterministic pivots  |  Auditable wave rules  |  Causal scoring  |  Multi-market scanning")
r.italic = True
r.font.size = Pt(10.5)
r.font.color.rgb = RGBColor.from_string(MUTED)
add_picture(
    doc,
    ROOT / "terminal-overview.png",
    Inches(6.45),
    "Elliott Wave Terminal showing market controls, status cards, chart, and Structure Inspector.",
).alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(14)
r = p.add_run("Prepared for the current 25-market test build")
r.bold = True
r.font.color.rgb = RGBColor.from_string(GREEN)
p = doc.add_paragraph("Version 1.0  |  July 2026")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

page_break(doc)
doc.add_heading("Purpose and safety boundary", level=1)
doc.add_paragraph(
    "The terminal converts canonical five-minute OHLCV data into volatility-adjusted pivots, "
    "tests deterministic Elliott structures, ranks surviving alternatives, and exposes exact "
    "invalidation and target references. It is designed to help a trader organize evidence. "
    "It does not predict with certainty and it does not replace position sizing or execution judgment."
)
add_note(
    doc,
    "Critical",
    "A Confidence Score is a deterministic guideline score, not a probability of profit. "
    "Never read 90/100 as a 90% win rate.",
    RED,
)
doc.add_heading("What the system does", level=2)
for text in (
    "Builds all chart timeframes dynamically from raw M5 candles.",
    "Extracts pivots with a rolling ATR threshold instead of a fixed percentage.",
    "Prunes candidates immediately when an absolute Elliott rule fails.",
    "Tracks alternate candidate paths, structure stage, lifecycle, invalidation, and targets.",
    "Scores Fibonacci fit, momentum, alternation, channeling, and pattern-specific geometry.",
    "Scans every installed market and registered timeframe for active opportunities.",
):
    bullet(doc, text)
doc.add_heading("What it deliberately does not do", level=2)
for text in (
    "It does not place orders or connect to a broker in the current build.",
    "It does not use machine learning, visual guessing, or subjective hand-drawn counts.",
    "It does not guarantee data completeness when a temporary provider omits candles.",
    "It does not remove the need to check liquidity, spread, news, and portfolio exposure.",
):
    bullet(doc, text)
doc.add_heading("Guide map", level=2)
add_table(
    doc,
    ["Section", "Use it when you need to..."],
    [
        ("Quick start", "Reach a defensible chart interpretation in a few minutes."),
        ("Terminal controls", "Understand every selector and slider."),
        ("Structures and stages", "Interpret Impulse, ZigZag, Flat, Triangle, Forming, and EntryReady."),
        ("Chart and inspector", "Read overlays, RSI, score audit, invalidation, target, and hints."),
        ("Global scanner", "Rank active setups across all markets and timeframes."),
        ("Operating checklist", "Apply consistent pre-trade and post-trade discipline."),
        ("Troubleshooting", "Resolve empty charts, missing arrows, or incomplete macro bars."),
    ],
    [2200, 7160],
)

page_break(doc)
doc.add_heading("Quick start: a disciplined five-minute workflow", level=1)
step(doc, "Select the market", "Choose the asset database in the Market list. Confirm DATA THROUGH is current enough for your decision.")
step(doc, "Select the timeframe", "Start with 4H or 1D for context, then inspect 1H or 30M for timing. Do not mix labels from separate timeframes as if they were one count.")
step(doc, "Keep default pivot settings initially", "Use ATR Multiplier 2.0 and ATR Period 14 until you have completed formal sensitivity testing for the asset.")
step(doc, "Choose Balanced and Actionable", "Balanced shows both motive and corrective structures; Actionable removes historical structures that no longer matter.")
step(doc, "Inspect the focused path", "Read Stage, Lifecycle, Confidence Score, invalidation, target zone, and System Hints before looking at the decorative shape.")
step(doc, "Open Score audit", "Confirm where the points came from. A high total built from weak or irrelevant evidence should be challenged.")
step(doc, "Check the confidence gate", "An arrow is shown only when the path is enabled, tradeable, and at or above the selected alert threshold.")
step(doc, "Apply independent risk controls", "Check reward-to-risk after friction, size the trade, and identify event risk before acting.")
add_note(
    doc,
    "Recommended sequence",
    "Context timeframe -> execution timeframe -> structure validity -> confidence evidence -> "
    "invalidation -> target -> friction-adjusted risk/reward -> position size.",
    GREEN,
)

doc.add_heading("Traffic-light interpretation", level=2)
add_table(
    doc,
    ["Terminal state", "Meaning", "Operator response"],
    [
        ("Forming", "An early partial path is developing.", "Watch only; do not treat the final pivot as confirmed."),
        ("EntryReady + below threshold", "The terminal pivot exists, but guideline quality is insufficient.", "Wait or reject. No marker should appear."),
        ("EntryReady + passed threshold", "Structure and confidence gates pass.", "Perform risk, execution, and market-regime checks."),
        ("Completed", "The full historical structure is confirmed.", "Use for context; it may no longer be an entry."),
        ("Invalidated / target hit", "The active trade thesis is no longer live.", "Do not enter from the old signal."),
    ],
    [1850, 3300, 4210],
)

page_break(doc)
doc.add_heading("Single Chart Terminal: control reference", level=1)
add_picture(
    doc,
    ROOT / "terminal-overview.png",
    Inches(6.7),
    "Single Chart Terminal with asset, timeframe, pattern and ATR controls above the chart.",
).alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph("Figure 1. Current terminal overview. Controls appear above the chart; the Structure Inspector is on the right.")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].italic = True
p.runs[0].font.size = Pt(8.5)
p.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

page_break(doc)
add_table(
    doc,
    ["Control", "Options / range", "How to use it"],
    [
        ("Market", "All local .db/.sqlite assets", "Selects one canonical M5 database. The current build includes 25 test markets."),
        ("Timeframe", "15M, 30M, 1H, 2H, 4H, 1D, 1W", "Resampled dynamically. Lower frames are noisier; higher frames need more complete history."),
        ("Pattern View", "Balanced; Impulse 1-5; ZigZag ABC", "Balanced is the default review. Filter when the chart is crowded or your thesis is pattern-specific."),
        ("ATR Multiplier", "1.5 to 4.0", "Higher values require larger reversals and produce fewer pivots. Lower values detect more detail and more noise."),
        ("ATR Period", "5 to 50", "Controls volatility smoothing. Fourteen is the standard baseline."),
        ("Candidate Scope", "Actionable; Recent 30D; All history", "Actionable is best for decisions. All history is mainly for audit and research."),
    ],
    [1650, 2350, 5360],
)
add_note(
    doc,
    "Parameter discipline",
    "Do not move ATR settings until a preferred pattern appears. That is outcome-driven fitting. "
    "Define and test a parameter policy before reviewing the setup.",
    GOLD,
)

doc.add_heading("How ATR settings change the chart", level=2)
add_table(
    doc,
    ["Change", "Expected effect", "Main risk"],
    [
        ("Lower multiplier", "More pivots and more short-duration paths", "Noise and over-counting"),
        ("Higher multiplier", "Fewer, larger swings", "Late recognition and missed smaller structures"),
        ("Shorter ATR period", "Threshold reacts faster to volatility", "Unstable behavior after isolated shocks"),
        ("Longer ATR period", "Smoother threshold", "Slow adaptation when volatility regime changes"),
    ],
    [1900, 3560, 3900],
)

page_break(doc)
doc.add_heading("Reading the terminal from top to bottom", level=1)
doc.add_heading("Market summary cards", level=2)
for text in (
    "LAST shows the most recent completed candle close and its one-bar percentage change.",
    "TIMEFRAME confirms the active aggregation. Always verify it before interpreting labels.",
    "VALID PATHS counts candidates after the current Pattern View and Candidate Scope filters.",
    "DATA THROUGH states the timestamp of the final completed bar, in UTC.",
):
    bullet(doc, text)
doc.add_heading("Legend and chart overlays", level=2)
add_table(
    doc,
    ["Visual", "Interpretation"],
    [
        ("Solid connected wave line", "Completed candidate structure."),
        ("Dashed wave line", "Forming/provisional structure. Its terminal state is not fully confirmed."),
        ("Green upward arrow", "Bullish EntryReady setup that passes the active confidence threshold."),
        ("Red downward arrow", "Bearish EntryReady setup that passes the active confidence threshold."),
        ("Start, 1-5, A-E labels", "Exact pivot coordinates used by the candidate; these are engine outputs, not manual drawings."),
        ("RSI pane with 30/70 guides", "Causal Wilder RSI used for momentum scoring and divergence context."),
        ("Red invalidation line", "Price boundary that breaks the active structural thesis."),
        ("Green target zone lines", "Deterministic Fibonacci projection references, not guaranteed exits."),
    ],
    [2600, 6760],
)
doc.add_heading("Marker threshold, Buy at, Sell at", level=2)
doc.add_paragraph(
    "This strip explains why a chart arrow is present or absent. The terminal evaluates the enabled top alternatives. "
    "A marker requires all three gates: the overlay is enabled, the candidate ends at a tradeable Wave 4 or Wave B, "
    "and its Confidence Score is at least the Alert Confidence Score."
)
add_note(
    doc,
    "Example",
    "If the strip says 'Below threshold (62.5)' while the threshold is 75, the structure may be EntryReady but no arrow "
    "is supposed to appear. Lowering the threshold would change visibility, not structural validity.",
    BLUE,
)

page_break(doc)
doc.add_heading("Structure Inspector", level=2)
add_table(
    doc,
    ["Panel", "What to verify"],
    [
        ("Focused path", "Rank, pattern, stage, direction, and score of the selected top alternative."),
        ("Alternative checkboxes", "Enable only the paths you want drawn and eligible for markers."),
        ("Selected structure", "Pattern family, direction, stage, and current lifecycle."),
        ("Confidence Score", "Guideline quality subject to maturity caps; never a probability."),
        ("Floor/Ceiling invalidation", "Exact structural boundary. Floor applies to bullish; ceiling applies to bearish."),
        ("Fibonacci target zone", "Projected price interval used for planning and scanner comparison."),
        ("System Hints", "Plain-language gate status, next required event, and trading interpretation."),
        ("Focus chart", "Zooms the visible history around the selected path; turn off for broader context."),
    ],
    [2450, 6910],
)

doc.add_heading("Patterns supported by the engine", level=1)
doc.add_heading("Impulse: Start-1-2-3-4-5", level=2)
doc.add_paragraph(
    "The standard impulse is directional and must satisfy hard structural rules. Wave 2 cannot retrace beyond the "
    "origin of Wave 1; Wave 3 must move beyond Wave 1 and cannot be the shortest motive wave; Wave 4 cannot fully "
    "retrace Wave 3 or enter Wave 1 territory; the standard policy requires Wave 5 to progress beyond Wave 3."
)
add_note(doc, "Use", "An EntryReady impulse ends at Wave 4 and scouts the potential Wave 5 move.", GREEN)
doc.add_heading("ZigZag: Start-A-B-C", level=2)
doc.add_paragraph(
    "A ZigZag is a sharp correction. Wave B remains below the Flat threshold and Wave C must travel in the correction "
    "direction beyond the Wave A extreme. An EntryReady ZigZag ends at Wave B and scouts Wave C."
)
doc.add_heading("Flat: Start-A-B-C", level=2)
add_table(
    doc,
    ["Subtype", "Deterministic classification"],
    [
        ("Regular Flat", "Wave B retraces approximately 90%-100% of A; C reaches or exceeds the A extreme."),
        ("Expanded Flat", "Wave B exceeds the A origin, within the configured extension bound; C reaches or exceeds A."),
        ("Running Flat", "Wave B exceeds the A origin while C fails the A extreme but meets minimum progress."),
    ],
    [2100, 7260],
)
doc.add_heading("Triangle: Start-A-B-C-D-E", level=2)
doc.add_paragraph(
    "Triangles are five-leg corrective structures. The engine checks containment, AC/BD boundary geometry, contraction "
    "or barrier behavior, terminal E placement, and boundary convergence. It rejects expanding geometry under the "
    "current standard triangle policy."
)
add_note(
    doc,
    "Hard rule versus guideline",
    "A hard-rule failure deletes the candidate. Fibonacci, RSI, alternation, and channeling can only change its score; "
    "they cannot rescue an invalid structure.",
    RED,
)

doc.add_heading("Stages and lifecycle are different", level=2)
add_table(
    doc,
    ["Field", "Question answered", "Typical values"],
    [
        ("Stage", "How much of the wave sequence has the engine observed?", "Forming, EntryReady, Completed"),
        ("Lifecycle", "Is the setup still relevant at the latest price?", "Forming, Active, Target hit, Invalidated"),
    ],
    [1500, 4800, 3060],
)

page_break(doc)
doc.add_heading("Confidence Score and score audit", level=1)
doc.add_paragraph(
    "The score is an auditable ranking out of 100. Completed candidates may use the full scale. Provisional candidates "
    "are maturity-capped so an incomplete structure cannot look as certain as a completed one."
)
add_table(
    doc,
    ["Category", "Maximum", "Evidence considered"],
    [
        ("Fibonacci alignment", "50", "Wave retracements, extensions, and pattern-specific ratios."),
        ("Momentum verification", "30", "Causal Wilder RSI strength and Wave 3/Wave 5 momentum behavior."),
        ("Channeling / alternation", "20", "Corrective depth/style contrast, duration symmetry, and geometry."),
    ],
    [2500, 1300, 5560],
)
doc.add_heading("How to inspect a score", level=2)
step(doc, "Open Score audit", "Read every ScoreItem reason and earned/maximum value.")
step(doc, "Identify concentration", "A high total dominated by one category is less balanced than a score supported across categories.")
step(doc, "Check stage cap", "Forming candidates are intentionally capped below fully mature candidates.")
step(doc, "Compare alternatives", "A small score difference does not erase structural differences or risk/reward differences.")
step(doc, "Retain the audit trail", "Record the selected path, score breakdown, threshold, invalidation, and target at decision time.")
add_note(
    doc,
    "Threshold guidance",
    "The default 75 threshold is an operational filter, not a universal optimum. Change it only after walk-forward and "
    "sensitivity analysis, then keep it stable during live evaluation.",
    GOLD,
)

doc.add_heading("Interpreting System Hints", level=2)
add_table(
    doc,
    ["Hint row", "Meaning"],
    [
        ("Pattern state", "Pattern, direction, and terminal wave represented by the selected candidate."),
        ("Confidence gate", "Pass/fail comparison against the active threshold."),
        ("Lifecycle", "Whether current price still leaves the thesis active."),
        ("Entry gate", "Whether a tradeable terminal Wave 4 or Wave B is present."),
        ("Marker decision", "Exact reason an arrow is shown or hidden."),
        ("Next required event", "What must occur before the setup advances."),
        ("Trading interpretation", "Concise decision-support summary, not an order instruction."),
        ("Invalidation reference", "Price and boundary direction that breaks the structure."),
        ("Target zone", "Current Fibonacci projection interval."),
    ],
    [2350, 7010],
)

page_break(doc)
doc.add_heading("Global Market Scanner", level=1)
add_picture(
    doc,
    ROOT / "scanner-overview.png",
    Inches(6.7),
    "Global Market Scanner table ranked by Confidence Score across 25 markets and seven timeframes.",
).alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph("Figure 2. Scanner results after evaluating 25 markets across seven registered timeframes.")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].italic = True
p.runs[0].font.size = Pt(8.5)
p.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
doc.add_paragraph(
    "The scanner runs the same pivot, engine, and scoring pipeline used by the chart. It includes Forming watch "
    "candidates and EntryReady Wave 4/B setups, filters out candidates whose lifecycle is no longer active, and sorts "
    "the result by Confidence Score."
)
doc.add_heading("Scanner workflow", level=2)
step(doc, "Set ATR parameters in the terminal", "The scanner uses the active ATR period and multiplier.")
step(doc, "Open Global Market Scanner", "Confirm the universe line shows the expected market and timeframe counts.")
step(doc, "Run market scan", "Wait for the master table and active-setup count.")
step(doc, "Prioritize EntryReady", "Forming rows are watch candidates; EntryReady rows have the tradeable terminal pivot.")
step(doc, "Compare direction and timeframe", "Avoid treating correlated rows as independent opportunities.")
step(doc, "Open the asset in Single Chart Terminal", "Verify the path visually and inspect the full score audit and lifecycle.")
step(doc, "Export if needed", "Use the table CSV control for an offline review record.")

doc.add_heading("Scanner columns", level=2)
add_table(
    doc,
    ["Column", "Interpretation"],
    [
        ("Market", "Local asset database name."),
        ("Timeframe", "Aggregation on which the candidate was detected."),
        ("Pattern", "Impulse, ZigZag, Flat, or Triangle."),
        ("Setup Stage", "Forming or EntryReady."),
        ("Direction", "Expected direction of the next modeled leg."),
        ("Confidence Score", "Deterministic rank after maturity cap."),
        ("Invalidation Price", "Hard structural boundary."),
        ("Fibonacci Target Zone", "Projected planning range."),
    ],
    [2050, 7310],
)
add_note(
    doc,
    "Correlation",
    "Several rows can represent the same underlying exposure: for example NIFTY, NIFTY Bank, HDFC Bank, and ICICI Bank. "
    "Count portfolio risk, not scanner rows.",
    RED,
)

page_break(doc)
doc.add_heading("Alerts and operations", level=1)
doc.add_paragraph(
    "Open Alerts & operations below the terminal. Set the Telegram Chat ID and Alert Confidence Score. The bot token "
    "belongs in .streamlit/secrets.toml as TELEGRAM_BOT_TOKEN; it is intentionally not entered in the browser."
)
add_note(
    doc,
    "Current limitation",
    "The current webhook shell logs a clean alert message but does not send a network request. Treat it as an integration "
    "template until a live provider and secure notification service are connected.",
    GOLD,
)
doc.add_heading("When an alert is eligible", level=2)
for text in (
    "The candidate is an active, tradeable Wave 4 or Wave B setup.",
    "Its Confidence Score meets or exceeds the selected threshold.",
    "The candidate has not already produced the same alert key in the current session.",
):
    bullet(doc, text)

doc.add_heading("Data freshness and refresh procedure", level=1)
doc.add_paragraph(
    "All analysis reads provider-neutral SQLite databases containing canonical M5 candles. Higher timeframes are generated "
    "on demand, which avoids synchronization drift. The current test universe can be refreshed from Yahoo Finance with:"
)
p = doc.add_paragraph()
p.style = "Intense Quote"
r = p.add_run("python market_data.py --period 60d")
r.font.name = "Consolas"
r.font.size = Pt(10)
doc.add_paragraph(
    "Yahoo Finance is a temporary, credential-free test source. Intraday retention and candle completeness are limited. "
    "Production use should replace only the adapter, not the database or analytical contracts."
)
add_table(
    doc,
    ["Profile", "Examples", "Calendar handling"],
    [
        ("NSE", "NIFTY, Indian equities", "Asia/Kolkata 09:15-15:30 sessions"),
        ("US equity", "S&P 500, Nasdaq 100", "America/New_York 09:30-16:00 sessions"),
        ("24/7", "BTC, ETH, SOL", "288 M5 candles per complete day"),
        ("FX", "EUR/USD, GBP/USD, USD/INR", "Weekday profile; provider gaps may remove macro bars"),
        ("Futures", "Gold, Silver, WTI", "Daily maintenance break accounted for"),
    ],
    [1800, 2800, 4760],
)

page_break(doc)
doc.add_heading("Pre-trade operating checklist", level=1)
doc.add_heading("Structure", level=2)
for text in (
    "The selected path is EntryReady, not merely Forming.",
    "Lifecycle is Active; target has not already been hit and invalidation has not occurred.",
    "The final pivot is Wave 4 for an impulse or Wave B for a corrective setup.",
    "The selected alternative remains visible and its labels match the intended candidate.",
):
    bullet(doc, text)
doc.add_heading("Evidence", level=2)
for text in (
    "Confidence passes the pre-defined threshold.",
    "Score audit shows acceptable Fibonacci, momentum, and geometry evidence.",
    "Higher-timeframe direction does not create an unexplained conflict.",
    "ATR parameters match the tested policy and were not tuned after seeing the pattern.",
):
    bullet(doc, text)
doc.add_heading("Risk and execution", level=2)
for text in (
    "Entry, invalidation, and target are on the correct side for the direction.",
    "Reward-to-risk remains acceptable after spread, slippage, and commission.",
    "Position size is derived from the invalidation distance, not from desired profit.",
    "Open positions, correlated exposure, and daily loss limits remain within policy.",
    "Upcoming earnings, macro releases, market closure, and liquidity conditions are checked.",
):
    bullet(doc, text)
add_note(
    doc,
    "No-trade rule",
    "If any hard structural, data-quality, lifecycle, or risk gate is uncertain, the correct terminal action is no trade.",
    RED,
)

doc.add_heading("Post-decision audit record", level=2)
add_table(
    doc,
    ["Record", "Minimum value to retain"],
    [
        ("Market context", "Asset, timeframe, timestamp, data-through timestamp"),
        ("Candidate", "Pattern, variant, direction, stage, node sequence"),
        ("Decision evidence", "Score total and ScoreItem breakdown"),
        ("Risk references", "Entry reference, invalidation, target zone, size, friction"),
        ("Outcome", "Filled/not filled, exit reason, realized P/L, maximum adverse excursion"),
    ],
    [2450, 6910],
)

page_break(doc)
doc.add_heading("Troubleshooting", level=1)
add_table(
    doc,
    ["Symptom", "Likely explanation", "Action"],
    [
        ("No valid paths", "Current pivots do not satisfy hard rules or the scope/filter removes them.", "Use Balanced + All history for diagnosis; do not weaken rules just to create a count."),
        ("No chart arrow", "Overlay disabled, setup not EntryReady, lifecycle inactive, or score below threshold.", "Read Marker threshold and System Hints; verify all three marker gates."),
        ("No complete candles", "Provider history is insufficient for that market/timeframe.", "Try a lower timeframe or refresh data. Never fabricate macro bars."),
        ("Weekly FX/futures absent", "Strict constituent M5 count was not met.", "Use complete lower frames or a higher-quality production feed."),
        ("Too many pivots", "ATR multiplier is too low for current volatility.", "Return to the tested baseline; evaluate changes through sensitivity analysis."),
        ("Too few pivots", "ATR multiplier is high or history is short.", "Check data coverage before changing parameters."),
        ("Different alternatives overlap", "The DAG found multiple valid paths through nearby pivots.", "Focus one path and compare exact labels, invalidation, stage, and score."),
        ("Scanner row differs from chart", "ATR settings, candidate scope, data refresh, or lifecycle changed.", "Match parameters and rerun the scanner immediately before inspection."),
        ("Dashboard shows old markets", "The Streamlit process or browser session is stale.", "Restart Streamlit and hard-refresh the browser."),
    ],
    [1800, 3650, 3910],
)
doc.add_heading("Command reference", level=2)
add_table(
    doc,
    ["Task", "Command"],
    [
        ("Start dashboard", "python -m streamlit run app.py"),
        ("Refresh test data", "python market_data.py --period 60d"),
        ("Run all tests", "python -m pytest -q"),
    ],
    [2500, 6860],
)

doc.add_heading("Glossary", level=1)
add_table(
    doc,
    ["Term", "Definition"],
    [
        ("ATR", "Average True Range; the volatility scale used by the pivot threshold."),
        ("Pivot", "A sequentially confirmed swing high or low with timestamp, price, type, and ATR."),
        ("Candidate", "An immutable labeled path through ordered pivots."),
        ("DAG", "Directed acyclic graph used to represent alternate sequential wave paths."),
        ("Hard rule", "A structural requirement whose failure deletes the candidate."),
        ("Guideline", "Fibonacci, momentum, or geometry evidence that changes score only."),
        ("Invalidation", "Exact price boundary that breaks the active structural thesis."),
        ("Target zone", "A Fibonacci projection interval used as a planning reference."),
        ("Forming", "Provisional structure that has not reached the tradeable terminal pivot."),
        ("EntryReady", "Partial structure ending at a completed Wave 4 or Wave B."),
        ("Lifecycle", "Current relevance of the candidate at the latest available price."),
        ("Confidence Score", "Auditable guideline ranking out of 100; not a probability."),
    ],
    [1950, 7410],
)

doc.add_heading("Final operating principle", level=1)
add_note(
    doc,
    "Use the terminal as a falsification engine",
    "Start by asking what price, rule, data condition, or risk limit would make the thesis wrong. "
    "Only then consider targets. The system is most valuable when it prevents weak trades, not when it decorates charts.",
    GREEN,
)

core = doc.core_properties
core.title = "Elliott Wave Terminal User Guide"
core.subject = "Operator manual for the deterministic Elliott Wave decision-support dashboard"
core.author = "Elliott Wave DSS Project"
core.keywords = "Elliott Wave, decision support, user guide, trading, Streamlit"
doc.save(OUT)
print(OUT)
